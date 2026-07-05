import json
from flask import Response, stream_with_context
import os
import time
import json
from flask import Flask, render_template, request, jsonify, send_from_directory
from docx import Document
from agents import Agent

app = Flask(__name__)

CONVERSATIONS_FILE = "conversations.json"

# ============ SPECIALIST AGENTS ============

researcher = Agent(
    name="Researcher",
    system_prompt="You are a research specialist. You gather and explain factual information clearly and concisely. When relevant, use web search to find current, accurate information.",
    provider="google",
    use_web_search=True
)

writer = Agent(
    name="Writer",
    system_prompt="You are a writing specialist. You take information and turn it into clear, well-organized, polished text.",
    provider="anthropic"
)

comms = Agent(
    name="Comms",
    system_prompt="""You are a business communications specialist.
You draft professional client emails, meeting recaps, and workspace messages.
You write in a warm, clear, professional tone appropriate for small business clients.""",
    provider="anthropic"
)

planner = Agent(
    name="Planner",
    system_prompt="You are a planning specialist. You break big goals or tasks into clear, numbered, actionable steps.",
    provider="anthropic"
)

coder = Agent(
    name="Coder",
    system_prompt="You are a coding specialist. You write clean, working code and explain it simply, assuming the reader is a beginner unless told otherwise.",
    provider="anthropic"
)

pm = Agent(
    name="ProjectManager",
    system_prompt="""You are a project manager for a small AI development team.
You break client requests into clear milestones, timelines, and task lists.
You track what's been delivered vs. what's pending, and write concise status updates
suitable for sending to a client or reviewing internally.""",
    provider="google"
)

sales = Agent(
    name="Sales",
    system_prompt="""You are a sales and proposal specialist for a freelance AI development business.
You write clear, persuasive client proposals, cost estimates, and outreach emails.
You know typical freelance pricing ($300-$1,500 for small business tools, $20-$100/month maintenance).
You write confidently but never overpromise technical capabilities.""",
    provider="openai"
)

recruiter = Agent(
    name="Recruiter",
    system_prompt="""You are a recruiting specialist for a small freelance tech team.
You write clear job posts, screening questions, and interview questions for roles like
developers, designers, and virtual assistants. You help define what skills are actually
needed for a given task before recommending hiring anyone.""",
    provider="openai"
)

infra = Agent(
    name="Infra",
    system_prompt="""You are an infrastructure and deployment specialist.
You give clear, practical recommendations on hosting, domains, deployment, and tooling
for small client projects, favoring low-cost, low-maintenance solutions appropriate for
a solo freelancer or small team, not enterprise-scale complexity.""",
    provider="openai"
)

agents = {
    "researcher": researcher,
    "writer": writer,
    "planner": planner,
    "coder": coder,
    "sales": sales,
    "pm": pm,
    "recruiter": recruiter,
    "infra": infra,
    "comms": comms,
}

# ============ ROUTER ============

def router(user_task):
    routing_prompt = f"""A user gave this task: "{user_task}"

Decide which specialist agent(s) should handle this, IN ORDER, comma-separated.
- researcher — facts, information, current events, explanations
- writer — turning info into polished written text
- planner — breaking a task into steps
- coder — writing or explaining code
- sales — proposals, pricing, client pitches
- pm — project milestones, timelines, status tracking
- recruiter — job posts, hiring criteria, interview questions
- infra — hosting, deployment, technical setup recommendations
- comms — client emails, meeting notes, professional messages

Example: "Write a proposal and break it into milestones" -> sales, pm

Reply with ONLY the agent name(s), comma-separated, nothing else."""

    decision = researcher.run(routing_prompt).strip().lower()
    chain = [name.strip() for name in decision.split(",") if name.strip() in agents]
    return chain if chain else ["researcher"]


# ============ DOCX GENERATION ============

def text_to_docx(text, filepath):
    doc = Document()
    for raw_line in text.split("\n"):
        line = raw_line.rstrip()

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.strip().startswith(("- ", "* ")):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
        else:
            p = doc.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True

    doc.save(filepath)


def save_agent_output(agent_name, text):
    os.makedirs("outputs", exist_ok=True)
    existing = len([f for f in os.listdir("outputs") if f.startswith(agent_name)])
    filename = f"{agent_name}_{existing + 1}.docx"
    filepath = os.path.join("outputs", filename)
    text_to_docx(text, filepath)
    return filename


# ============ CONVERSATION LOG (explicit, reliable) ============

def load_conversations():
    if os.path.exists(CONVERSATIONS_FILE):
        with open(CONVERSATIONS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def save_conversation(user_message, steps):
    conversations = load_conversations()
    conversations.append({
        "id": len(conversations),
        "title": user_message[:60] + ("..." if len(user_message) > 60 else ""),
        "ts": time.time(),
        "user_message": user_message,
        "steps": steps
    })
    with open(CONVERSATIONS_FILE, "w", encoding="utf-8") as f:
        json.dump(conversations, f, indent=2)


# ============ ROUTES ============

@app.route("/")
def home():
    return render_template("index.html")


@app.route("/world")
def world():
    return render_template("world.html")


@app.route("/chat", methods=["POST"])
def chat():
    user_message = request.json.get("message", "")

    def generate():
        chain = router(user_message)
        context = user_message
        steps = []

        for agent_name in chain:
            agent = agents[agent_name]
            result = agent.run(context)
            filename = save_agent_output(agent_name, result)
            step = {"agent": agent_name, "text": result, "download": filename}
            steps.append(step)
            context = f"Original task: {user_message}\n\nPrevious agent's output:\n{result}\n\nContinue the task."
            yield json.dumps(step) + "\n"

        save_conversation(user_message, steps)

    return Response(stream_with_context(generate()), mimetype="application/x-ndjson")


@app.route("/download/<filename>")
def download(filename):
    return send_from_directory("outputs", filename, as_attachment=True)


@app.route("/memory", methods=["GET"])
def memory():
    summaries = [agents[name].memory_summary() for name in agents]
    return jsonify({"agents": summaries})


@app.route("/memory/clear/<agent_name>", methods=["POST"])
def clear_memory(agent_name):
    if agent_name in agents:
        agents[agent_name].clear_memory()
        return jsonify({"status": "cleared", "agent": agent_name})
    return jsonify({"status": "not found"}), 404


@app.route("/conversations", methods=["GET"])
def conversations():
    convos = load_conversations()
    convos.reverse()  # most recent first
    return jsonify({"conversations": convos})


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)